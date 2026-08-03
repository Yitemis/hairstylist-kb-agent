# -*- coding: utf-8 -*-
"""Word（docx）解析器。

基于 python-docx 开源库。
借鉴思路：保留标题层级（来自九阳 POC 章节定位调优）。
"""
from __future__ import annotations

import logging
from typing import List

from app.rag.parsers.doc_types import ChildChunk, ElementType, ParentChunk
from app.rag.parsers.utils import download_file, is_safe_url
from app.rag.chunkers.smart_chunker import build_child_chunks, build_parent_chunks

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 200 * 1024 * 1024


class DocxParser:
    """Word 文档解析器。

    关键：保留 Heading 1/2/3 样式（章节定位的根基）。
    """

    def __init__(self, file_uri: str, filename: str):
        if is_safe_url(file_uri):
            self.file_url = file_uri
            self.file_path = None
        else:
            self.file_path = file_uri
            self.file_url = None
        self.filename = filename

    def load(
        self,
        document_id: str = "",
        tenant_id: str = "default",
        category: str = "general",
        chunk_size: int = 800,
        chunk_overlap: int = 80,
        parent_chunk_size: int = 2000,
    ) -> List[ParentChunk]:
        text = self._read_text()
        # 转成 Markdown 风格（保留 ## 标题）
        md_text = self._convert_to_markdown(text)
        from app.rag.chunkers.smart_chunker import (
            extract_qa_pairs,
            merge_qa_into_chunks,
            split_markdown_by_heading,
        )
        sections = split_markdown_by_heading(md_text, chunk_size, chunk_overlap)
        qa_pairs = extract_qa_pairs(md_text)
        if qa_pairs:
            sections = merge_qa_into_chunks(sections, qa_pairs)
        child_chunks = build_child_chunks(
            sections,
            source_filename=self.filename,
            document_id=document_id,
            tenant_id=tenant_id,
            category=category,
        )
        return build_parent_chunks(
            child_chunks,
            parent_chunk_size=parent_chunk_size,
            source_filename=self.filename,
            document_id=document_id,
            tenant_id=tenant_id,
        )

    def _read_text(self) -> str:
        try:
            from docx import Document
            import io
            binary = self._read_file()
            doc = Document(io.BytesIO(binary))
            return self._extract_with_styles(doc)
        except ImportError:
            logger.warning("python-docx 未安装，回退到简单文本提取")
            return self._read_text_fallback()

    def _extract_with_styles(self, doc) -> str:
        """提取带样式信息的文本（保留 Heading 层级）。"""
        lines = []
        for para in doc.paragraphs:
            style_name = (para.style.name or "").lower() if para.style else ""
            text = para.text.strip()
            if not text:
                continue
            if "heading 1" in style_name or "标题 1" in style_name:
                lines.append(f"# {text}")
            elif "heading 2" in style_name or "标题 2" in style_name:
                lines.append(f"## {text}")
            elif "heading 3" in style_name or "标题 3" in style_name:
                lines.append(f"### {text}")
            else:
                lines.append(text)
        # 提取表格（HTML）
        from docx.oxml.ns import qn
        for tbl in doc.tables:
            html = self._table_to_html(tbl)
            lines.append(html)
        return "\n".join(lines)

    def _table_to_html(self, tbl) -> str:
        """把 Word 表格转成 HTML（保留列关系）。"""
        rows = []
        for row in tbl.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("<tr>" + "".join(f"<td>{c}</td>" for c in cells) + "</tr>")
        return "<table>" + "".join(rows) + "</table>"

    def _convert_to_markdown(self, text: str) -> str:
        """已经是 markdown 风格了（_extract_with_styles 输出 # 标题），直接返回。"""
        return text

    def _read_text_fallback(self) -> str:
        """降级方案：用二进制解码。"""
        binary = self._read_file()
        return binary.decode("utf-8", errors="ignore")

    def _read_file(self) -> bytes:
        if self.file_url:
            return download_file(self.file_url, MAX_FILE_SIZE)
        with open(self.file_path, "rb") as f:
            return f.read()
